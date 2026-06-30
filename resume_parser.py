import re
import os
import json
import PyPDF2
import docx
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.corpus import stopwords

# Download required NLTK data (only first run)
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    nltk.download('punkt', quiet=True)
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords', quiet=True)
try:
    nltk.data.find('tokenizers/punkt_tab')
except LookupError:
    nltk.download('punkt_tab', quiet=True)

# ── Common resume sections ──────────────────────────────────────────────────
SECTION_KEYWORDS = {
    'education': ['education', 'academic', 'qualification', 'degree', 'university', 'college', 'school'],
    'experience': ['experience', 'employment', 'work history', 'professional experience', 'career'],
    'skills': ['skills', 'technical skills', 'core competencies', 'technologies', 'proficiencies'],
    'projects': ['projects', 'personal projects', 'academic projects', 'key projects'],
    'certifications': ['certifications', 'certificates', 'courses', 'training', 'licenses'],
    'summary': ['summary', 'objective', 'profile', 'about', 'overview'],
    'achievements': ['achievements', 'awards', 'honors', 'accomplishments'],
    'contact': ['contact', 'email', 'phone', 'address', 'linkedin', 'github'],
}

# ── Common tech skills to detect ────────────────────────────────────────────
TECH_SKILLS = [
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'swift', 'kotlin',
    'react', 'angular', 'vue', 'node.js', 'django', 'flask', 'fastapi', 'spring', 'express',
    'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'sqlite', 'oracle',
    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'terraform', 'ansible',
    'git', 'github', 'gitlab', 'bitbucket', 'jenkins', 'ci/cd',
    'pandas', 'numpy', 'matplotlib', 'scikit-learn', 'tensorflow', 'pytorch', 'keras',
    'html', 'css', 'bootstrap', 'tailwind', 'sass',
    'linux', 'unix', 'bash', 'powershell',
    'rest', 'api', 'graphql', 'microservices', 'agile', 'scrum',
    'excel', 'tableau', 'power bi', 'data analysis', 'machine learning',
    'selenium', 'pytest', 'junit', 'postman',
]


def extract_text_from_pdf(filepath):
    """Extract plain text from a PDF file."""
    text = ''
    try:
        with open(filepath, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
    except Exception as e:
        print(f"PDF extraction error: {e}")
    return text.strip()


def extract_text_from_docx(filepath):
    """Extract plain text from a DOCX file."""
    text = ''
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            text += para.text + '\n'
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + ' '
                text += '\n'
    except Exception as e:
        print(f"DOCX extraction error: {e}")
    return text.strip()


def extract_resume_text(filepath):
    """Detect file type and extract text."""
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.docx':
        return extract_text_from_docx(filepath)
    return ''


def extract_email(text):
    pattern = r'[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}'
    matches = re.findall(pattern, text)
    return matches[0] if matches else None


def extract_phone(text):
    pattern = r'(\+?\d[\d\s\-().]{7,}\d)'
    matches = re.findall(pattern, text)
    return matches[0].strip() if matches else None


def extract_linkedin(text):
    pattern = r'linkedin\.com/in/[\w\-]+'
    matches = re.findall(pattern, text, re.IGNORECASE)
    return matches[0] if matches else None


def extract_github(text):
    pattern = r'github\.com/[\w\-]+'
    matches = re.findall(pattern, text, re.IGNORECASE)
    return matches[0] if matches else None


def detect_sections(text):
    """Detect which standard resume sections are present."""
    found = []
    text_lower = text.lower()
    for section, keywords in SECTION_KEYWORDS.items():
        if any(kw in text_lower for kw in keywords):
            found.append(section)
    return found


def extract_skills(text):
    """Find known tech skills mentioned in the resume."""
    text_lower = text.lower()
    found = []
    for skill in TECH_SKILLS:
        if skill in text_lower:
            found.append(skill)
    return list(set(found))


def get_word_count(text):
    words = re.findall(r'\b\w+\b', text)
    return len(words)


def parse_resume(filepath):
    """
    Full resume parse.
    Returns a dict with extracted metadata.
    """
    text = extract_resume_text(filepath)
    if not text:
        return {'error': 'Could not extract text from file', 'text': ''}

    result = {
        'text': text,
        'email': extract_email(text),
        'phone': extract_phone(text),
        'linkedin': extract_linkedin(text),
        'github': extract_github(text),
        'sections_found': detect_sections(text),
        'extracted_skills': extract_skills(text),
        'word_count': get_word_count(text),
    }
    return result
